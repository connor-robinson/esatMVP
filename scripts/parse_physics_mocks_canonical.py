# -*- coding: utf-8 -*-
"""Parse ESAT Physics Mock Modules DOCX into canonical JSON for import."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

DOCX = Path(r"c:\Users\anson\Downloads\ESAT_Physics_Mock_Modules_A_B.docx")
OUT_DIR = Path(r"c:\Users\anson\Desktop\nocalcMVP2_real\tmp_physics_mocks")
OUT_DIR.mkdir(parents=True, exist_ok=True)

DIFF_LABEL = {
    "1/4": "Foundation",
    "2/4": "Standard",
    "3/4": "Challenging",
    "4/4": "Very challenging",
}


def cell_text(cell) -> str:
    return cell.text.replace("\r\n", "\n").replace("\r", "\n").strip()


def nested_tables(table: Table) -> list[Table]:
    el = table._element
    return [Table(child, table._parent) for child in el.findall(".//" + qn("w:tbl")) if child is not el]


def parse_options(table: Table) -> dict[str, str]:
    options: dict[str, str] = {}
    for nt in nested_tables(table):
        # option tables are 6x2 with letter | text
        if len(nt.rows) < 4 or len(nt.columns) < 2:
            continue
        first_left = cell_text(nt.rows[0].cells[0])
        if first_left in ("Choice", "A", "B", "C", "D", "E", "F", "G", "H") or re.match(r"^[A-H]$", first_left):
            if first_left == "Choice":
                continue  # distractor header table; skip here
            for row in nt.rows:
                letter = cell_text(row.cells[0]).strip()
                text = cell_text(row.cells[1]).strip()
                if re.match(r"^[A-H]$", letter) and text:
                    options[letter] = text
    return options


def parse_candidate(table: Table, number: int) -> dict:
    cell = table.rows[0].cells[0]
    paras = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    # paras[0] is "N.", paras[1] is stem (possibly multi)
    stem_parts = []
    for i, p in enumerate(paras):
        if i == 0 and re.match(r"^\d+\.?$", p):
            continue
        if re.match(r"^\d+\.\s*", p):
            stem_parts.append(re.sub(r"^\d+\.\s*", "", p))
        else:
            stem_parts.append(p)
    options = parse_options(table)
    has_diagram = ("drawing" in table._element.xml) or ("blip" in table._element.xml)
    return {
        "number": number,
        "stem": "\n".join(stem_parts).strip(),
        "options": options,
        "hasDiagram": has_diagram,
    }


def parse_distractors(table: Table) -> dict[str, str]:
    distractors: dict[str, str] = {}
    for nt in nested_tables(table):
        if len(nt.rows) < 2 or len(nt.columns) < 2:
            continue
        header0 = cell_text(nt.rows[0].cells[0])
        header1 = cell_text(nt.rows[0].cells[1]) if len(nt.rows[0].cells) > 1 else ""
        if header0 == "Choice" or "wrong" in header1.lower():
            for row in nt.rows[1:]:
                letter = cell_text(row.cells[0]).strip()
                text = cell_text(row.cells[1]).strip()
                if re.match(r"^[A-H]$", letter) and text:
                    distractors[letter] = text
    return distractors


def parse_meta_header(table: Table) -> dict:
    """Parse nested header like 'A1 | P1.2 ... | Difficulty 2/4' and TARGET / EDITOR PICK."""
    meta = {
        "topicCode": None,
        "topicName": None,
        "difficultyFraction": None,
        "targetSeconds": None,
        "editorPick": False,
        "headerRaw": None,
    }
    for nt in nested_tables(table):
        if len(nt.rows) != 1:
            continue
        left = cell_text(nt.rows[0].cells[0])
        right = cell_text(nt.rows[0].cells[1]) if len(nt.rows[0].cells) > 1 else ""
        if "P" in left and "Difficulty" in left:
            meta["headerRaw"] = left + " || " + right
            # A1   |   P1.2 Electric circuits   |   Difficulty 2/4
            parts = [p.strip() for p in re.split(r"\s*\|\s*", left) if p.strip()]
            if len(parts) >= 2:
                topic_full = parts[1]
                tm = re.match(r"^(P\d+(?:\.\d+)?)\s+(.*)$", topic_full)
                if tm:
                    meta["topicCode"] = tm.group(1)
                    meta["topicName"] = tm.group(2).strip()
                else:
                    meta["topicName"] = topic_full
            for p in parts:
                dm = re.search(r"Difficulty\s+(\d/\d)", p)
                if dm:
                    meta["difficultyFraction"] = dm.group(1)
            if "EDITOR PICK" in right.upper() or "★" in right:
                meta["editorPick"] = True
            tm = re.search(r"TARGET\s+(\d+)\s*s", right, re.I)
            if tm:
                meta["targetSeconds"] = int(tm.group(1))
    return meta


def parse_editor(table: Table, number: int, module: str, overview_row: dict) -> dict:
    paras = []
    # Get paragraph texts from outermost cell only (not nested), in order
    # The outer table is usually 1 cell containing paras + nested tables
    outer_cell = table.rows[0].cells[0]
    for p in outer_cell.paragraphs:
        t = p.text.strip()
        if t:
            paras.append(t)

    def field_after(prefix: str) -> str | None:
        for p in paras:
            if p.startswith(prefix):
                return p[len(prefix) :].strip()
        return None

    answer_raw = field_after("Answer")
    tip = field_after("Tip")
    solution = field_after("Solution")
    benchmark = field_after("Benchmark note")
    if benchmark is None:
        benchmark = field_after("Benchmark")

    difficulty_line = None
    for p in paras:
        if p.startswith("Difficulty:"):
            difficulty_line = p
            break

    difficulty_word = None
    target_from_line = None
    if difficulty_line:
        # Difficulty: Standard   |   Suggested target: 70 seconds
        m = re.search(r"Difficulty:\s*([^|]+)", difficulty_line)
        if m:
            difficulty_word = m.group(1).strip()
        m = re.search(r"Suggested target:\s*(\d+)\s*seconds?", difficulty_line, re.I)
        if m:
            target_from_line = int(m.group(1))

    answer_letter = None
    answer_text = None
    if answer_raw:
        am = re.match(r"^([A-H])\s+(.*)$", answer_raw.strip(), re.S)
        if am:
            answer_letter = am.group(1)
            answer_text = am.group(2).strip()

    meta = parse_meta_header(table)
    distractors = parse_distractors(table)

    frac = overview_row.get("difficulty") or meta.get("difficultyFraction")
    label = DIFF_LABEL.get(frac or "", difficulty_word or "")
    # Exact difficulty string preferred as "2/4 Standard" style matching doc scale
    if frac and label:
        difficulty = f"{frac} {label}"
    elif frac:
        difficulty = frac
    else:
        difficulty = difficulty_word

    target_seconds = (
        overview_row.get("targetSeconds")
        or meta.get("targetSeconds")
        or target_from_line
    )
    # Prefer exact overview target display "70 s"
    target_display = overview_row.get("targetDisplay")
    if not target_display and target_seconds:
        target_display = f"{target_seconds} s"

    topic_code = overview_row.get("topicCode") or meta.get("topicCode")
    topic_name = overview_row.get("topicName") or meta.get("topicName")
    # overview topic is like "P1.2 Electric circuits"
    if not topic_code and overview_row.get("topicFull"):
        tm = re.match(r"^(P\d+(?:\.\d+)?)\s+(.*)$", overview_row["topicFull"])
        if tm:
            topic_code = tm.group(1)
            topic_name = tm.group(2)

    editor_pick = meta.get("editorPick") or False
    # overview may mark gold - but we rely on ★ EDITOR PICK in header

    return {
        "module": module,
        "number": number,
        "editorPick": editor_pick,
        "answer": answer_letter,
        "answerText": answer_text,
        "topicCode": topic_code,
        "topicName": topic_name,
        "difficulty": difficulty,
        "targetSeconds": target_seconds,
        "targetDisplay": target_display,
        "tip": tip,
        "solution": solution,
        "distractors": distractors,
        "benchmarkNote": benchmark,
        "difficultyWord": difficulty_word,
        "difficultyLine": difficulty_line,
        "metaHeader": meta.get("headerRaw"),
        "paras": paras,
    }


def parse_overview(table: Table) -> list[dict]:
    rows = []
    for ri, row in enumerate(table.rows):
        if ri == 0:
            continue
        cells = [cell_text(c) for c in row.cells]
        # columns: Question, Answer, Topic, Difficulty, Target
        q = cells[0].strip()
        m = re.match(r"^([AB])?(\d+)$", q)
        if not m:
            continue
        topic_full = cells[2].strip()
        tm = re.match(r"^(P\d+(?:\.\d+)?)\s+(.*)$", topic_full)
        target_display = cells[4].strip()
        ts = None
        tsm = re.match(r"^(\d+)\s*s", target_display)
        if tsm:
            ts = int(tsm.group(1))
        rows.append(
            {
                "label": q,
                "module": m.group(1),
                "number": int(m.group(2)),
                "answer": cells[1].strip(),
                "topicFull": topic_full,
                "topicCode": tm.group(1) if tm else None,
                "topicName": tm.group(2).strip() if tm else topic_full,
                "difficulty": cells[3].strip(),
                "targetDisplay": target_display,
                "targetSeconds": ts,
            }
        )
    return rows


def map_diagram_images(doc: Document) -> dict[str, str]:
    """Map question keys like A7 to image filenames by document order of drawings in candidate tables."""
    # Candidate A tables 8-34, B 38-64
    mapping = {}
    image_files = sorted((OUT_DIR / "images").glob("image*.png"), key=lambda p: int(re.search(r"(\d+)", p.stem).group(1)))
    # Actually images in zip may not be sequential with questions - need rId order in doc
    return mapping


def extract_question_image_map(doc: Document) -> dict[str, str]:
    """Walk candidate question tables and associate embedded images in order of appearance."""
    # Build relationship id -> media name from package
    import zipfile

    rid_to_media = {}
    with zipfile.ZipFile(DOCX) as z:
        # parse document.xml.rels
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        for m in re.finditer(r'Id="(rId\d+)"[^>]*Target="media/([^"]+)"', rels):
            rid_to_media[m.group(1)] = m.group(2)
        # also TargetMode
        for m in re.finditer(r'Id="(rId\d+)"[^>]*Target="media/([^"]+)"', rels):
            pass

    def images_in_table(table: Table) -> list[str]:
        xml = table._element.xml
        found = []
        for m in re.finditer(r'r:embed="(rId\d+)"', xml):
            rid = m.group(1)
            if rid in rid_to_media:
                found.append(rid_to_media[rid])
        return found

    result = {}
    for i, num in enumerate(range(1, 28)):
        imgs = images_in_table(doc.tables[8 + i])
        if imgs:
            result[f"A{num}"] = imgs[0]
    for i, num in enumerate(range(1, 28)):
        imgs = images_in_table(doc.tables[38 + i])
        if imgs:
            result[f"B{num}"] = imgs[0]
    return result


def main() -> None:
    doc = Document(str(DOCX))
    overview_a = parse_overview(doc.tables[67])
    overview_b = parse_overview(doc.tables[95])
    ov_a_by_n = {r["number"]: r for r in overview_a}
    ov_b_by_n = {r["number"]: r for r in overview_b}

    module_a = []
    module_b = []

    for i in range(27):
        num = i + 1
        cand = parse_candidate(doc.tables[8 + i], num)
        ed = parse_editor(doc.tables[68 + i], num, "A", ov_a_by_n[num])
        # Consistency: answer letter from overview
        assert ed["answer"] == ov_a_by_n[num]["answer"], (num, ed["answer"], ov_a_by_n[num]["answer"])
        assert ed["answer"] in cand["options"], (num, ed["answer"], cand["options"].keys())
        # answer text should match option
        if ed["answerText"] != cand["options"][ed["answer"]]:
            # allow minor whitespace
            if ed["answerText"].replace(" ", "") != cand["options"][ed["answer"]].replace(" ", ""):
                print("WARN A answer text mismatch", num, repr(ed["answerText"]), repr(cand["options"][ed["answer"]]))
        # distractors: all except correct
        expected = set(cand["options"]) - {ed["answer"]}
        if set(ed["distractors"]) != expected:
            print("WARN A distractor keys", num, sorted(ed["distractors"]), sorted(expected))
        module_a.append({**cand, **ed, "id": f"A{num}"})

    for i in range(27):
        num = i + 1
        cand = parse_candidate(doc.tables[38 + i], num)
        ed = parse_editor(doc.tables[96 + i], num, "B", ov_b_by_n[num])
        assert ed["answer"] == ov_b_by_n[num]["answer"], (num, ed["answer"], ov_b_by_n[num]["answer"])
        assert ed["answer"] in cand["options"], (num, ed["answer"], cand["options"].keys())
        if ed["answerText"] != cand["options"][ed["answer"]]:
            if ed["answerText"].replace(" ", "") != cand["options"][ed["answer"]].replace(" ", ""):
                print("WARN B answer text mismatch", num, repr(ed["answerText"]), repr(cand["options"][ed["answer"]]))
        expected = set(cand["options"]) - {ed["answer"]}
        if set(ed["distractors"]) != expected:
            print("WARN B distractor keys", num, sorted(ed["distractors"]), sorted(expected))
        module_b.append({**cand, **ed, "id": f"B{num}"})

    image_map = extract_question_image_map(doc)

    out = {
        "moduleA": module_a,
        "moduleB": module_b,
        "imageMap": image_map,
        "meta": {
            "timeMinutes": 40,
            "questionsPerModule": 27,
            "calculator": "Not permitted",
            "coverage": {
                "P1": {"name": "Electricity", "A": 5, "B": 4},
                "P2": {"name": "Magnetism", "A": 3, "B": 4},
                "P3": {"name": "Mechanics", "A": 7, "B": 7},
                "P4": {"name": "Thermal physics", "A": 3, "B": 3},
                "P5": {"name": "Matter", "A": 3, "B": 3},
                "P6": {"name": "Waves", "A": 4, "B": 4},
                "P7": {"name": "Radioactivity", "A": 2, "B": 2},
            },
        },
    }
    path = OUT_DIR / "canonical.json"
    path.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    print("wrote", path)
    print("A opts sample", module_a[0]["options"])
    print("A1", module_a[0]["answer"], module_a[0]["topicCode"], module_a[0]["difficulty"], module_a[0]["targetSeconds"])
    print("A7 editorPick", module_a[6]["editorPick"], "diagram", module_a[6]["hasDiagram"])
    print("imageMap", image_map)
    print(
        "counts",
        len(module_a),
        len(module_b),
        sum(1 for q in module_a if q["hasDiagram"]),
        sum(1 for q in module_b if q["hasDiagram"]),
    )


if __name__ == "__main__":
    main()
